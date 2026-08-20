# -*- coding: utf-8 -*-
import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_rtl(table):
    tblPr = table._tbl.tblPr
    tblCellMar = parse_xml(f'<w:bidiVisual {nsdecls("w")} w:val="1"/>')
    tblPr.append(tblCellMar)

def add_heading_rtl(doc, text, level=1, color_rgb=(15, 23, 42)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.bidi = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Cairo'
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color_rgb)
    return p

def add_paragraph_rtl(doc, text, bold=False, color_rgb=(51, 65, 85), font_size=10.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.bidi = True
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Cairo'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color_rgb)
    return p

categories = [
    {
        "cat_num": 1,
        "title": "المحور الأول: الذكاء الاصطناعي والأتمتة السريرية والصيدلانية (Clinical & Pharmacy AI)",
        "color": "0F766E",
        "ideas": [
            ("1", "فاحص التداخلات الدوائية التفاعلي اللحظي (Drug-Drug Interactions Analyzer)", "تحذير فوري عند إضافة صنفين متعارضين في فاتورة واحدة مع توثيق BNF 83.", "عالية 🔥", "Next.js Modal + BNF 83 Rules Engine"),
            ("2", "حاسبة جرعات الأطفال التلقائية بناءً على الوزن والعمر (Pediatric Dosing Calculator)", "حساب جرعات المضادات والمسكنات التلقائي بالسنتيمتر/مل لمنع الأخطاء الطبية.", "عالية 🔥", "Custom JS Engine + BNF Monograph Doses"),
            ("3", "المساعد الصيدلاني الصوتي للبحث عن البدائل (Voice AI Pharmacy Assistant)", "البحث الصوتي عن بدائل الأدوية بنفس التركيبة والمادة الفعالة وسعرها.", "متوسطة ⚡", "Web Speech API + Semantic DB Search"),
            ("4", "قارئ الروشتات الطبية الذكي بالرؤية الحاسوبية (Prescription OCR & Vision AI)", "تصوير الروشتة المكتوبة بخط اليد وتحويلها تلقائياً لأصناف وسلة مشتريات.", "مستقبلية 🚀", "Gemini 1.5 Flash Vision / Cloud OCR"),
            ("5", "فاحص سلامة الأدوية للحوامل والمرضعات (Pregnancy & Lactation Safety Check)", "إظهار تصنيف FDA/BNF لخطورة الدواء للحوامل والمرضعات مع بدائل آمنة.", "عالية 🔥", "PostgreSQL Tags + BNF Knowledge DB"),
            ("6", "حاسبة التعديل الكلوي لجرعات الأدوية (Renal Clearance Dosage Adjuster)", "حساب معدل GFR وتعديل جرعة المضادات لمرضى الكلى تلقائياً.", "متوسطة ⚡", "Cockcroft-Gault Equation Component"),
            ("7", "منشئ خطط العلاج المزمن وتذكير المرضى (Chronic Medication Auto-Refill)", "تتبع مرضى السكري والضغط وإرسال تذكير واتساب لتجديد الدواء قبل نفاده.", "عالية 🔥", "Cron Job + n8n WhatsApp Dispatcher"),
            ("8", "محرك التعرف على الأدوية بالصورة من شريط الدواء (Pill Identification by Image)", "التعرف على الصنف وتاريخ صلاحيته عبر التقاط صورة للشريط أو العلبة.", "مستقبلية 🚀", "Mobile Camera + AI Image Embedding"),
            ("9", "نظام توثيق وحساب التركيبات الصيدلانية (Compounding Master Formulation)", "حساب نسب المواد الخام وتكلفة التصنيع وتوليد ملصق الجرعة والتشغيلة.", "متوسطة ⚡", "Formula Builder Module + Label Printer"),
            ("10", "كاشف التشابه الإملائي والصوتي للأدوية (LASA - Look-Alike Sound-Alike Alert)", "تنبيه الصيدلي عند اختيار دواء يشبه دواءً آخر خطيراً لمنع صرف الدواء الخاطئ.", "عالية 🔥", "Levenshtein String Distance Algorithm")
        ]
    },
    {
        "cat_num": 2,
        "title": "المحور الثاني: القياسات الحيوية وإثبات الحضور المتقدم (Biometrics & AI Geofencing)",
        "color": "1E40AF",
        "ideas": [
            ("11", "التحقق من الوجه بالذكاء الاصطناعي مع كشف الحيوية (AI Face Liveness Verification)", "التقاط صورة سريعة عند الدخول والتأكد من وجود الموظف الحقيقي لمنع التحايل.", "عالية 🔥", "Face-api.js / MediaPipe Liveness Detection"),
            ("12", "نطاق جغرافي متعدد المقرات والفروع (Multi-Location Smart Geofencing)", "إمكانية تحديد عدة فروع للمنشأة مع انتقال الصلاحية للموظف المتنقل تلقائياً.", "عالية 🔥", "Haversine Distance Multi-Coordinates Array"),
            ("13", "الحضور عبر شبكة الواي فاي المعتمدة للفرع (Office Wi-Fi BSSID Match)", "مطابقة معرّف شبكة الواي فاي للفرع للتأكد من وجود الهاتف داخل المنشأة حصراً.", "متوسطة ⚡", "Network API / Captive Portal Handshake"),
            ("14", "بطاقة الحضور التفاعلية الديناميكية عبر NFC (NFC Mobile Tap Check-In)", "تسجيل الحضور بلمس الهاتف لملصق NFC مشفر ومثبت في الاستقبال.", "متوسطة ⚡", "Web NFC API + Rotating Nonce Tokens"),
            ("15", "رمز الاستجابة السريعة المتغير لحظياً (Dynamic Rotating QR Kiosk)", "شاشة لوحية في المدخل تعرض QR متغير كل 10 ثوانٍ يمسحه الموظف بهاتفه.", "عالية 🔥", "WebSocket / Server-Sent Events (SSE)"),
            ("16", "حظر أجهزة المحاكاة ومواقع الـ GPS المزيفة (Anti-Spoofing & Mock Location Guard)", "كشف ومنع محاولات تزييف الموقع الجغرافي وتنبيه الإدارة فوراً بمحاولة التحايل.", "عالية 🔥", "Navigator Geolocation Accuracy Validator"),
            ("17", "تسجيل الحضور الجماعي لشفتات الفرق (Team Lead Group Check-In)", "تمكين مشرف الورشة أو الوردية من تسجيل حضور فريقه بضغطة زر واحدة.", "متوسطة ⚡", "Batch Attendance API Endpoint"),
            ("18", "الحضور التلقائي الذكي عبر البلوتوث منخفض الطاقة (BLE Beacon Proximity)", "تسجيل الدخول تلقائياً بمجرد دخول الموظف نطاق جهاز البلوتوث (Beacon).", "مستقبلية 🚀", "Web Bluetooth API / Mobile PWA"),
            ("19", "التحقق الصوتي من الموظف (Voice Biometrics Authentication)", "نطق الموظف لكود عشوائي للتحقق من بصمة صوته عند الحضور عن بُعد.", "مستقبلية 🚀", "Audio Context Waveform Matcher"),
            ("20", "تتبع مسار المهام الميدانية خطوة بخطوة (Field Breadcrumb GPS Route)", "رسم خريطة مسار الفني الميداني ومواقع الزيارات ومطابقتها مع المهام.", "عالية 🔥", "Leaflet Map Path Tracking + Polyline View")
        ]
    },
    {
        "cat_num": 3,
        "title": "المحور الثالث: إدارة الشفتات والورديات المعقدة والرواتب (Smart Shifts & Payroll)",
        "color": "047857",
        "ideas": [
            ("21", "حاسبة الرواتب والضرائب والخصومات المتقدمة بنقرة واحدة (One-Click Payroll Run)", "توليد كشوفات الرواتب الشاملة للبدلات والمكافآت والسلف بـ (د.ل) وتصديرها PDF/Excel.", "عالية 🔥", "Dual Salary Engine + PDF Export Lib"),
            ("22", "سوق تبادل الشفتات بين الموظفين (Shift Swap Marketplace & Approvals)", "طلب الموظف تبادل ورديته مع زميل، مع وصول إشعار للمدير للموافقة الفورية.", "عالية 🔥", "Shift Request Workflow + WhatsApp Action"),
            ("23", "محرك الجدولة الذكية للورديات وتغطية النواقص (Auto-Rostering AI)", "توزيع الشفتات تلقائياً بناءً على كفاءة الموظفين والحد الأدنى للكوادر لكل قسم.", "متوسطة ⚡", "Constraint Satisfaction Scheduling Algorithm"),
            ("24", "حساب مكافآت الإنتاجية ومبيعات الأصناف (Sales Commission Tracker)", "حساب نسبة العمولة لكل صيدلي أو فني بناءً على إنجازاته ومبيعاته الموثقة.", "عالية 🔥", "Sales Target DB + Commission Matrix"),
            ("25", "نظام طلب ومتابعة السلف والعهد المالية (Advance Salary & Petty Cash)", "تقديم طلب سلفة من الهاتف وخصمها المجدول تلقائياً من راتب الشهر المستحق.", "عالية 🔥", "Loan Management Sub-Schema + Ledger"),
            ("26", "رصيد الإجازات السنوية والمرضية وسجل الغياب (Vacation Balance & Leave Flow)", "احتساب رصيد الإجازات المستحقة وتطبيق الخصومات وسجلات التقرير الطبي.", "عالية 🔥", "Leave Accrual System + Document Upload"),
            ("27", "حساب ساعات العمل الإضافي التلقائي (Tiered Overtime Multiplier)", "احتساب الساعات بعد نهاية الدوام بمعدل 1.5x أو 2.0x في العطلات الرسمية.", "عالية 🔥", "Rate Rules Engine Multiplier Integration"),
            ("28", "محفظة الموظف الرقمية وكشف الحساب التفاعلي (Employee Digital Payslip Wallet)", "واجهة هاتف للموظف تعرض تفاصيل أرباحه وساعاته وسلفه بشفافية لحظية.", "عالية 🔥", "Next.js Mobile-Optimized Dashboard Tab"),
            ("29", "نظام التقييم الدوري وتأثيره على سلم الرواتب (KPI Appraisal & Bonus System)", "تقييم شهري للموظف من قبل مديره يؤثر على نسبة الحوافز السنوية.", "متوسطة ⚡", "Scoring Matrix + Performance Analytics"),
            ("30", "تنبيهات تجاوز الحد الأقصى لساعات الدوام الأسبوعية (Labor Fatigue Guard)", "تنبيه الإدارة عند اقتراب الموظف من إجهاد العمل لتجنب الإرهاق والأخطاء.", "متوسطة ⚡", "Weekly Hours Aggregator & Alert Trigger")
        ]
    },
    {
        "cat_num": 4,
        "title": "المحور الرابع: إدارة العمليات الميدانية وتذاكر الصيانة (Field Service & Operations)",
        "color": "B45309",
        "ideas": [
            ("31", "توقيع العميل الإلكتروني الموثق بـ OTP و GPS (Cryptographic Client e-Signature)", "توقيع العميل على شاشة الهاتف مع رمز تحقق OTP وموقع الزيارة الجغرافي.", "عالية 🔥", "HTML5 Canvas + SHA-256 Hash + WhatsApp OTP"),
            ("32", "توثيق صور الصيانة قبل وبعد الإنجاز (Before & After Photo Proof)", "إلزام الفني بالتقاط صور المعدة قبل الصيانة وبعدها وحفظها في التقرير.", "عالية 🔥", "Direct Image Upload + Cloudinary/Local Store"),
            ("33", "تسعير قطع الغيار الميدانية وسحبها من المخزن فوراً (Field Parts Deduction)", "اختيار قطع الغيار المستخدمة أثناء الزيارة وخصمها من المخزن وإضافتها للفاتورة.", "عالية 🔥", "Inventory Linkage + Realtime Stock Sync"),
            ("34", "إرسال رابط تتبع الفني على الخريطة للعميل (Uber-like Technician Live Tracking)", "رابط يصل للعميل على واتساب يوضح تحرك الفني ووصوله التقديري لموقع العمل.", "متوسطة ⚡", "Live Geolocation Stream + Public Map Token"),
            ("35", "تقييم العميل الفوري للخدمة بالنجوم والتعليق (Post-Visit Star Rating Survey)", "رسالة واتساب تلقائية للعميل فور اكتمال الزيارة لتقييم أداء الفني وجودة الخدمة.", "عالية 🔥", "WhatsApp Webhook Handler + CSAT Analytics"),
            ("36", "عقود الصيانة الدورية المجدولة تلقائياً (Recurring Maintenance Contracts - SLA)", "توليد زيارات صيانة دورية كل شهر أو 3 أشهر وإسنادها للمهندسين تلقائياً.", "عالية 🔥", "SLA Cron Dispatcher + Contract Model"),
            ("37", "توجيه الفنيين عبر خرائط جوجل بأقصر مسار (Smart Multi-Stop Route Optimizer)", "ترتيب زيارات اليوم للفنيين جغرافياً لتقليل استهلاك الوقود وزمن الوصول.", "متوسطة ⚡", "Google Maps / OSRM Routing Algorithm"),
            ("38", "الوضع غير المتصل بالإنترنت للزيارات الميدانية (Offline-First PWA Mode)", "إتمام الزيارات في الأماكن ضعيفة التغطية ومزامنتها تلقائياً فور عودة النت.", "عالية 🔥", "IndexedDB Storage + Background Sync Worker"),
            ("39", "إصدار إيصال سداد وبوليصة ضريبية فورية (Instant Thermal Bluetooth Receipt)", "طباعة الفاتورة والإيصال مباشرة عبر طابعات البلوتوث المحمولة لدى الفني.", "متوسطة ⚡", "Web Bluetooth Thermal Printer Command (ESC/POS)"),
            ("40", "نظام الضمان وإعادة فتح الزيارة المجانية (Warranty Tracking & Free Recall)", "تتبع فترة ضمان الصيانة وربط الزيارة بالسابقة إذا تكررت نفس المشكلة.", "عالية 🔥", "Warranty Expiry Logic + Free Callback Ticket")
        ]
    },
    {
        "cat_num": 5,
        "title": "المحور الخامس: أتمتة واتساب والبوتات التفاعلية (WhatsApp Intelligence & WAHA)",
        "color": "15803D",
        "ideas": [
            ("41", "بوت الحضور السريع عبر واتساب (WhatsApp One-Touch Attendance Bot)", "إرسال الموظف موقعه الجغرافي (Live Location) على واتساب لتسجيل الدخول.", "عالية 🔥", "WAHA Location Message Parser + Geofence API"),
            ("42", "بوت استعلام العملاء عن توفر الأدوية والأسعار (Medicine Availability Bot)", "يرسل العميل اسم الدواء أو الباركود ويرد البوت بالسعر والتوفر فوراً.", "عالية 🔥", "Inbound WhatsApp Webhook + Fuzzy DB Search"),
            ("43", "إشعارات وصول النواقص للمرضى المسجلين (Shortage Restock Alert Broadcast)", "إشعار واتساب تلقائي للعميل فور توفير الدواء الذي سبق وسأل عنه ولم يكن متوفراً.", "عالية 🔥", "Customer Waitlist Schema + Restock Trigger"),
            ("44", "التقرير المالي الصوتي اليومي للمدير (Daily Audio Executive Voice Digest)", "إرسال رسالة صوتية للمدير على واتساب بملخص المبيعات والحضور بصوت ذكي.", "متوسطة ⚡", "ElevenLabs / OpenAI TTS + n8n WhatsApp Media"),
            ("45", "بوت استلام طلبات وتوصيل الصيدلية (WhatsApp Pharmacy Order & Delivery Bot)", "استلام صورة الروشتة وتأكيد العنوان وإرسال رابط الدفع وسائق التوصيل.", "عالية 🔥", "Multi-Step WhatsApp State Machine"),
            ("46", "أزرار تفاعلية للموافقة على الإجازات (WhatsApp Interactive Approval Buttons)", "وصول طلب الإجازة للمدير بأزرار [موافقة] أو [رفض] مباشرة دون فتح الموقع.", "عالية 🔥", "WAHA Interactive Action Buttons Handler"),
            ("47", "تنبيهات انقطاع درجات حرارة ثلاجة الأدوية (IoT Vaccine Fridge Alarm)", "رسالة واتساب طارئة للمدير إذا تغيرت درجة حرارة ثلاجة الأنسولين واللقاحات.", "متوسطة ⚡", "ESP32 / MQTT Webhook + WhatsApp Urgent Alert"),
            ("48", "برودكاست العروض الطبية والتوعوية المخصصة (Targeted Broadcast Campaigns)", "إرسال رسائل توعوية وعروض للأمهات ومرضى السكري والعملاء المميزين.", "عالية 🔥", "Broadcast Audience Segmenter + Tenant Url"),
            ("49", "روبوت تدريب وصقل مهارات الموظفين اليومي (Daily 1-Minute Quiz Bot)", "إرسال سؤال صيدلاني أو تقني يومياً للموظفين لتطوير أدائهم ومكافأة الفائزين.", "متوسطة ⚡", "Micro-learning Question Bank + Scoreboard"),
            ("50", "إشعارات نفاد المخزون اللحظية للموردين (Automated Supplier Stock Alerts)", "توليد وإرسال أمر شراء PDF لمندوب المورد على واتساب فور وصول الحد الأدنى.", "عالية 🔥", "PDF Engine Stream + WAHA Document Dispatch")
        ]
    },
    {
        "cat_num": 6,
        "title": "المحور السادس: سلاسل الإمداد والمخزون الذكي (Supply Chain & Smart Inventory)",
        "color": "4338CA",
        "ideas": [
            ("51", "التنبؤ الذكي بالطلب ومعدل الاستهلاك (AI Demand & Seasonality Forecast)", "توقع الكميات المطلوبة في مواسم الشتاء والصيف لتجنب نقص الأدوية الموسمية.", "عالية 🔥", "Moving Average & Exponential Smoothing ML"),
            ("52", "تتبع تواريخ الصلاحية بطريقة FEFO المتقدمة (First-Expired, First-Out Tracking)", "إلزام البيع بالتشغيلة الأقرب انتهاءً مع تنبيه الصيدلي قبل 6 أشهر و 3 أشهر.", "عالية 🔥", "Batch Expiry Table + Color Coded Badges"),
            ("53", "مقارنة عروض أسعار الموردين الآلية (Supplier Price Quote Comparison Matrix)", "إظهار أفضل مورد يقدم أعلى نسبة بونص وأقل سعر للصنف المطلوب.", "عالية 🔥", "Supplier Purchase History + Best Price Ranker"),
            ("54", "الجرد السريع متعدد الكاميرات عبر الهواتف (Multi-Phone Concurrent Stocktaking)", "جرد الصيدلية أو المستودع بواسطة عدة موظفين في نفس اللحظة بكاميراتهم.", "عالية 🔥", "Concurrent Barcode Scanner + Realtime WS Count"),
            ("55", "تتبع إرجاع الأدوية للشركات والموردين (Supplier Return & Credit Note Tracker)", "إدارة الأصناف الوشيكة وإصدار إشعارات دائنة للموردين ومتابعة سدادها.", "عالية 🔥", "Return Vouchers Workflow + Accounting Link"),
            ("56", "الربط مع الفاتورة الإلكترونية والباركود الموحد (Unified National Drug Code Index)", "ربط الأصناف مع الدليل الدوائي الموحد وتوحيد الأسعار الرسمية والضرائب.", "متوسطة ⚡", "NDC Database Sync + Tax Compliant Engine"),
            ("57", "تتبع حركة الصنف وتحليل الركود (Dead Stock & Slow Moving Identification)", "اكتشاف الأصناف الراكدة التي لم تتحرك منذ 90 يوماً واقتراح عروض لتصريفها.", "عالية 🔥", "Inventory Velocity Metrics + Alert View"),
            ("58", "إدارة المستودعات الفرعية والتحويل بين الفروع (Inter-Branch Transfer Vouchers)", "طلب تحويل صنف من فرع لآخر مع توثيق الاستلام والتسليم برمز أمان.", "عالية 🔥", "Transfer Slip Model + 2-Step Verification"),
            ("59", "حساب هامش الربح الحقيقي بعد البونص والخصم (True Margin & Profit Analysis)", "احتساب التكلفة الفعلية للقطعة بعد احتساب بونص المورد والخصم النقدي.", "عالية 🔥", "Weighted Average Unit Cost Calculator"),
            ("60", "سجل التالف وفاقد التخزين وحوادث الكسر (Damage & Breakage Loss Tracker)", "توثيق الأدوية المنتهية أو التالفة مع الأسباب وتقرير خسائر دوري معتمد.", "متوسطة ⚡", "Write-Off Protocol + Financial Loss Report")
        ]
    },
    {
        "cat_num": 7,
        "title": "المحور السابع: منصة الاشتراكات وإدارة المستأجرين (Multi-Tenant SaaS Platform)",
        "color": "6D28D9",
        "ideas": [
            ("61", "بوابات الدفع الإلكتروني الليبية للاشتراكات (Sadad, Moamalat, Tadawul, T-Pay)", "تجديد اشتراك الأنشطة التجارية تلقائياً بالدينار الليبي ببطاقات الدفع الإلكتروني.", "عالية 🔥", "Libyan Payment Gateways Webhook Integrations"),
            ("62", "التسجيل الذاتي للشركات مع تجربة مجانية 14 يوماً (Self-Service Signup & Trial)", "إنشاء نشاط تجاري جديد ونطاق فرعي تلقائياً بدون تدخل بشري في 30 ثانية.", "عالية 🔥", "Automated Tenant Provisioning + Trial Cron"),
            ("63", "تخصيص الهوية البصرية والنطاق الخاص للعملاء (Custom Domain & White Labeling)", "ربط الشركة لنطاقها الخاص (مثل `hr.company.ly`) مع شعارها وألوانها الرسمية.", "عالية 🔥", "Reverse Proxy Wildcard + Dynamic CSS Variables"),
            ("64", "لوحة تحكم السوبر أدمن الشاملة لمراقبة السيرفرات (Super Admin Health & MRR)", "مراقبة الإيراد الشهري المتكرر (MRR) ونشاط السيرفر وحالة قواعد البيانات.", "عالية 🔥", "Live Metric Dashboard + Docker Health Checks"),
            ("65", "نظام الإحالة والعمولات للمسوقين (Affiliate & Partner Referral Program)", "توليد كود خصم للشركاء وحساب عمولاتهم على كل نشاط يشترك عبرهم.", "متوسطة ⚡", "Referral Tracking System + Payout Ledger"),
            ("66", "إيقاف وتجميد الخدمة التلقائي عند انتهاء الاشتراك (Grace Period & Auto-Suspension)", "إشعار النشاط قبل 7 أيام من نهاية الاشتراك وتجميد لوحة التحكم بلطف عند التعثر.", "عالية 🔥", "Subscription Expiry Middleware Guard"),
            ("67", "تصدير النسخ الاحتياطية للأنشطة بنقرة زر (Tenant Data Backup & Portability)", "تمكين كل شركة من تنزيل كامل بياناتها في ملف مضغوط (SQL/JSON/Excel).", "متوسطة ⚡", "Tenant-Isolated Dump Stream Generator"),
            ("68", "مركز مساعدة ودعم فني مدمج مع تذاكر فورية (In-App Support Ticket & Chat)", "فتح تذاكر دعم فني ومحادثة الدعم الفني مباشرة من لوحة التحكم.", "متوسطة ⚡", "Support Chat Component + Live Status"),
            ("69", "سجل التغييرات وإشعارات التحديثات للعملاء (In-App Release Notes & Roadmap)", "نافذة تنبثق للعملاء عند إضافة ميزات جديدة تشرح طريقة استخدامها بالفيديو.", "متوسطة ⚡", "What's New Modal + Changelog Markdown"),
            ("70", "تسعير مرن بحسب عدد الموظفين والفروع (Usage-Based Dynamic Billing)", "حساب سعر الاشتراك بدقة بحسب عدد الموظفين الفعليين النشطين شهرياً.", "عالية 🔥", "Dynamic Metered Billing Calculation Engine")
        ]
    },
    {
        "cat_num": 8,
        "title": "المحور الثامن: ذكاء الأعمال والتحليلات التنبؤية (BI & Predictive Analytics)",
        "color": "9D174D",
        "ideas": [
            ("71", "خريطة التركيز الحراري لحركة العملاء والذروة (Peak Hours & Heatmap Analysis)", "تحديد الساعات الأكثر ازدحاماً لتنظيم شفتات الموظفين وتوزيع الكوادر بكفاءة.", "عالية 🔥", "Recharts Hourly Traffic Visualizer"),
            ("72", "كاشف الشذوذ في المبيعات ومحاولات الاختلاس (AI Fraud & Anomaly Detection)", "تنبيه الإدارة عند وجود إلغاءات متكررة للفواتير أو فتح درج النقود دون بيع.", "عالية 🔥", "Statistical Outlier Detection Algorithm"),
            ("73", "تحليل كفاءة وإنتاجية الفنيين والمهندسين (Technician Performance Scorecard)", "مقارنة الفنيين من حيث سرعة حل المشاكل، التقييم، وتكلفة قطع الغيار.", "عالية 🔥", "Comparative Leaderboard & Radar Charts"),
            ("74", "مؤشر ولاء المرضى ومعدل العودة (Patient Retention & Churn Predictor)", "معرفة نسبة الزبائن الدائمين واكتشاف الذين توقفوا عن الشراء لإعادة استهدافهم.", "متوسطة ⚡", "Customer Frequency & Recency Matrix (RFM)"),
            ("75", "لوحة القيادة التنفيذية للتلفزيونات وشاشات الإدارة (Wallboard TV Kiosk Mode)", "واجهة أنيقة لعرض مؤشرات الأداء الحية على شاشات المكاتب وقاعات الاجتماعات.", "متوسطة ⚡", "Fullscreen Auto-Refreshing Wallboard View"),
            ("76", "تقارير أثر الطقس والمواسم على مبيعات الأدوية (Weather Impact Correlation)", "ربط مبيعات أدوية الحساسية والإنفلونزا بتغيرات الطقس ودرجات الحرارة.", "مستقبلية 🚀", "OpenWeather API Integration + Correlation ML"),
            ("77", "تصدير التقارير المجدول عبر البريد والواتساب (Automated Scheduled Reports)", "إرسال تقرير PDF أسبوعي فجر كل أحد إلى بريد مجلس الإدارة والمدراء.", "عالية 🔥", "Scheduled Cron + Puppeteer / PDF Renderer"),
            ("78", "حساب تكلفة الفرصة البديلة للنواقص (Lost Sales Opportunity Analysis)", "حساب إجمالي الأرباح الضائعة بسبب الأصناف غير المتوفرة التي طلبها الزبائن.", "عالية 🔥", "Shortage Search Count * Average Margin Formula"),
            ("79", "تحليل سلة المشتريات والبيع المتقاطع (Market Basket & Cross-Sell Engine)", "اقتراح أصناف مكملة (مثل فيتامين C مع المضاد الحيوي) لزيادة متوسط الفاتورة.", "عالية 🔥", "Apriori Association Rule Mining"),
            ("80", "مؤشر الانضباط العام والالتزام بالمواعيد (Company Punctuality Index - CPI)", "درجة تقييم عامة للمنشأة تقيس مدى التزام الكادر ككل باللوائح والورديات.", "متوسطة ⚡", "Weighted Scoring Aggregate Metric")
        ]
    },
    {
        "cat_num": 9,
        "title": "المحور التاسع: تجربة الموظف والتحفيز والألعاب (Gamification & Experience)",
        "color": "C2410C",
        "ideas": [
            ("81", "نظام أوسمة التميز وموظف الشهر (Employee of the Month & Badges)", "منح أوسمة رقمية تفاعلية (الأكثر انضباطاً، بطل المبيعات، منقذ الشفتات).", "عالية 🔥", "Gamified Profile Badges & Trophies Component"),
            ("82", "متجر مكافآت النقاط واستبدالها بإجازات وهدايا (Employee Rewards Store)", "جمع نقاط عن كل يوم التزام واستبدالها بساعات تأخير مسموحة أو جوائز عينية.", "متوسطة ⚡", "Points Ledger & Virtual Reward Redemption"),
            ("83", "بطاقة الموظف التفاعلية الذكية لمشاركتها عبر المحفظة (Apple & Google Wallet Pass)", "تصدير بطاقة الدوام الرقمية إلى Apple Wallet أو Google Wallet بالباركود.", "متوسطة ⚡", "PassKit (.pkpass) Generator Engine"),
            ("84", "حائط الإنجازات وتقدير الزملاء الداخلي (Kudos & Peer Recognition Wall)", "إمكانية إرسال بطاقات شكر وتقدير بين الزملاء داخل لوحة التحكم لتعزيز الروح.", "متوسطة ⚡", "Interactive Social Wall Component"),
            ("85", "استبيانات الرضا الوظيفي ومقترحات التطوير السرية (Anonymous Employee Feedback)", "صندوق مقترحات سري ومشفر للموظفين لإرسال ملاحظاتهم لإدارة المنشأة بأمان.", "عالية 🔥", "Encrypted Anonymous Submissions Queue"),
            ("86", "تذكير المناسبات وأعياد الميلاد والترقيات (Work Anniversary & Celebration Bot)", "تهنئة تلقائية للموظف بيوم ميلاده أو ذكرى انضمامه للعمل في مجموعة واتساب.", "متوسطة ⚡", "Celebration Trigger + Automated Group Graphic"),
            ("87", "الوضع المظلم المخصص عالي الفخامة (Ultra-Luxe Dark & OLED Mode)", "تصميم داكن موفر للطاقة ومريح للعين لصيادلة الشفتات الليلية والفنيين.", "عالية 🔥", "Tailwind Dark Theme Selector with Auto-Sunset"),
            ("88", "نظام دليل الموظف والتدريب التفاعلي (Interactive Onboarding Guide)", "دليل تفاعلي خطوة بخطوة للموظف الجديد لتعريفه بسياسات المكان واستخدام النظام.", "متوسطة ⚡", "Intro.js Step-by-Step Tour Flow"),
            ("89", "طلبات الزي الموحد والمعدات والأدوات (Uniform & Equipment Request Flow)", "طلب مقاسات المعاطف الطبية وأدوات الفحص واستلامها بتوقيع رقمي.", "متوسطة ⚡", "Assets Requisition Module"),
            ("90", "مؤشر اللياقة والنشاط داخل العمل (Workplace Steps & Wellness Challenge)", "تحدي ترفيهي لعدد الخطوات المقطوعة أثناء الدوام لتعزيز النشاط والحيوية.", "مستقبلية 🚀", "Pedometer Web API / Google Fit Sync")
        ]
    },
    {
        "cat_num": 10,
        "title": "المحور العاشر: الأمان المتقدم والامتثال والتطبيقات المدمجة (Enterprise Security & PWA)",
        "color": "334155",
        "ideas": [
            ("91", "تطبيق الويب التقدمي فائق السرعة مع التثبيت (Full Offline PWA + App Store Ready)", "تثبيت النظام كأيقونة تطبيق مباشر على شاشة الآيفون والأندرويد دون الحاجة لمتجر.", "عالية 🔥", "Service Worker + Manifest.json + Push API"),
            ("92", "المصادقة الثنائية ببصمة الإصبع ورمز الأمان (2FA / Passkeys / WebAuthn)", "دخول الإدارة عبر بصمة الإصبع أو FaceID للأجهزة الداعمة لأقصى درجات الحماية.", "عالية 🔥", "WebAuthn / Passkeys API Integration"),
            ("93", "سجل التدقيق الجنائي غير القابل للتعديل (Immutable Audit Log & Tamper-Proof)", "تسجيل كل حركة تعديل وقت أو حذف فاتورة مع الـ IP واسم المستخدم وتاريخ الثواني.", "عالية 🔥", "AuditLog PostgreSQL Model + Trigger Lock"),
            ("94", "عزل البيانات المشفر لكل نشاط تجاري (Row-Level Security & Encrypted Fields)", "تشفير أرقام الهواتف والبيانات الحساسة في قاعدة البيانات بمفاتيح AES-256.", "عالية 🔥", "Prisma Middleware Encryption / pgcrypto"),
            ("95", "كشف الدخول المتزامن من أجهزة متعددة غير مصرح بها (Device Fingerprint Guard)", "منع الموظف من تسجيل الحضور من هاتف زميله عبر قفل الحساب على بصمة الجهاز.", "عالية 🔥", "Device Fingerprinting (Canvas + UserAgent Hash)"),
            ("96", "النسخ الاحتياطي التلقائي المشفر سحابياً (Automated Encrypted DB Snapshots)", "نسخ احتياطي يومي مشفر يتم رفعه إلى سيرفرات S3 سحابية منعزلة للحماية من الكوارث.", "عالية 🔥", "Cron Script + pg_dump + GPG Encryption + S3"),
            ("97", "تحديد أوقات الوصول للوحة التحكم وحظر الدخول خارج الدوام (Access Time Window Guard)", "منع الموظف العادي من تصفح بيانات النظام خارج ساعات شفته المعتمدة.", "متوسطة ⚡", "Role-Based Shift Time Window Validator"),
            ("98", "كاشف الهجمات ومحاولات التخمين للـ PIN (Brute-Force Rate Limiter & IP Ban)", "قفل الحساب وحظر الـ IP تلقائياً بعد 3 محاولات إدخال رقم سري خاطئ.", "عالية 🔥", "Redis / In-Memory Rate Limiting Engine"),
            ("99", "الامتثال لخصوصية بيانات المرضى والملفات الطبية (HIPAA / GDPR Ready Compliance)", "إخفاء بيانات المريض الطبية وتشفيرها بما يوافق المعايير الدولية للسرية.", "عالية 🔥", "Data Anonymizer & Consent Management"),
            ("100", "وضع الطوارئ واستمرارية الأعمال (Disaster Recovery & Local Standalone Mode)", "تشغيل سيرفر محلي مصغر داخل الصيدلية يعمل حتى لو انقطع الإنترنت الخارجي بالكامل.", "عالية 🔥", "Local SQLite/Postgres Edge Node + Sync Agent")
        ]
    }
]

def generate_document():
    doc = docx.Document()
    
    # Page Setup (Margins)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Set RTL for Section Header & Footer
        sectPr = section._sectPr
        sectPr.append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))

    # Title Banner Table
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_table.autofit = False
    
    cell = title_table.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, "0F172A")
    set_cell_margins(cell, top=260, bottom=260, left=300, right=300)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.bidi = True
    r1 = p.add_run("🚀 خارطة الطريق والابتكار: 100 فكرة تفاعلية لمنظومة حضورك\n")
    r1.font.name = 'Cairo'
    r1.font.size = Pt(19)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(248, 250, 252)
    
    r2 = p.add_run("HodoorK Multi-Tenant SaaS & Pharmacy/Field Intelligence System\n")
    r2.font.name = 'Segoe UI'
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(148, 163, 184)
    
    r3 = p.add_run("دليل الميزات التنافسية والذكاء الاصطناعي والأتمتة الشاملة للأنشطة التجارية والطبية والخدمية")
    r3.font.name = 'Cairo'
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(56, 189, 248)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Executive Intro
    intro_p = add_paragraph_rtl(doc, "📋 مقدمة الدليل التفاعلي:", bold=True, font_size=12, color_rgb=(15, 23, 42))
    add_paragraph_rtl(doc, "تم إعداد هذه الوثيقة التفاعلية كمرجع هندسي واستثماري متقدم يحتوي على 100 ميزة وفكرة تطبيقية نوعية مصممة خصيصاً للتكامل مع البنية التحتية لمنظومة «حضورك» (Next.js 14, PostgreSQL, Prisma, WAHA/WhatsApp, BNF 83 Clinical AI, OCR, Geofencing). تم تقسيم الأفكار إلى 10 محاور رئيسية، مع تحديد آلية التنفيذ ومستوى الأولوية وخانات تفاعلية للتقييم والمتابعة.", font_size=10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Summary Statistics Table
    stats_table = doc.add_table(rows=2, cols=4)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_rtl(stats_table)
    
    headers = ["عدد الأفكار الإجمالي", "المحاور التقنية", "التقنيات المستخدمة", "العائد المتوقع (ROI)"]
    vals = ["100 فكرة نوعية", "10 محاور متكاملة", "Next.js / AI / WhatsApp / DB", "ريادة سوقية وتضاعف القيمة"]
    
    for i in range(4):
        c_head = stats_table.cell(0, i)
        c_val = stats_table.cell(1, i)
        set_cell_background(c_head, "1E293B")
        set_cell_background(c_val, "F1F5F9")
        set_cell_margins(c_head, 80, 80, 100, 100)
        set_cell_margins(c_val, 100, 100, 100, 100)
        
        p_h = c_head.paragraphs[0]
        p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_h.paragraph_format.bidi = True
        r = p_h.add_run(headers[i])
        r.font.name = 'Cairo'
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        
        p_v = c_val.paragraphs[0]
        p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_v.paragraph_format.bidi = True
        r_v = p_v.add_run(vals[i])
        r_v.font.name = 'Cairo'
        r_v.font.size = Pt(9.5)
        r_v.font.bold = True
        r_v.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Iterate through Categories
    for cat in categories:
        # Category Title Banner
        cat_p = doc.add_paragraph()
        cat_p.paragraph_format.bidi = True
        cat_p.paragraph_format.space_before = Pt(16)
        cat_p.paragraph_format.space_after = Pt(4)
        run_cat = cat_p.add_run(f"📂 {cat['title']}")
        run_cat.font.name = 'Cairo'
        run_cat.font.size = Pt(13)
        run_cat.font.bold = True
        
        # Color hex to RGB
        hex_c = cat['color']
        rgb_c = (int(hex_c[0:2], 16), int(hex_c[2:4], 16), int(hex_c[4:6], 16))
        run_cat.font.color.rgb = RGBColor(*rgb_c)

        # Table for Category
        table = doc.add_table(rows=1 + len(cat['ideas']), cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_rtl(table)
        
        col_widths = [Inches(0.4), Inches(2.2), Inches(2.3), Inches(0.9), Inches(1.1), Inches(0.5)]
        
        # Table Header
        col_names = ["#", "اسم الميزة والابتكار", "شرح الوظيفة والقيمة المضافة", "الأولوية", "البنية التقنية (Tech Stack)", "الحالة"]
        for j, name in enumerate(col_names):
            cell = table.cell(0, j)
            set_cell_background(cell, hex_c)
            set_cell_margins(cell, 90, 90, 80, 80)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.bidi = True
            r = p.add_run(name)
            r.font.name = 'Cairo'
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

        # Populate Rows
        for row_idx, idea in enumerate(cat['ideas'], start=1):
            num, title, desc, priority, tech = idea
            row_cells = [
                table.cell(row_idx, 0),
                table.cell(row_idx, 1),
                table.cell(row_idx, 2),
                table.cell(row_idx, 3),
                table.cell(row_idx, 4),
                table.cell(row_idx, 5)
            ]
            
            # Row alternating bg
            bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            for c in row_cells:
                set_cell_background(c, bg_color)
                set_cell_margins(c, 70, 70, 70, 70)
            
            # # Number
            p0 = row_cells[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p0.paragraph_format.bidi = True
            r0 = p0.add_run(num)
            r0.font.name = 'Cairo'
            r0.font.size = Pt(9)
            r0.font.bold = True
            r0.font.color.rgb = RGBColor(71, 85, 105)

            # Title
            p1 = row_cells[1].paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p1.paragraph_format.bidi = True
            r1 = p1.add_run(title)
            r1.font.name = 'Cairo'
            r1.font.size = Pt(9.5)
            r1.font.bold = True
            r1.font.color.rgb = RGBColor(15, 23, 42)

            # Description
            p2 = row_cells[2].paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p2.paragraph_format.bidi = True
            r2 = p2.add_run(desc)
            r2.font.name = 'Cairo'
            r2.font.size = Pt(8.5)
            r2.font.color.rgb = RGBColor(51, 65, 85)

            # Priority
            p3 = row_cells[3].paragraphs[0]
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.bidi = True
            r3 = p3.add_run(priority)
            r3.font.name = 'Cairo'
            r3.font.size = Pt(8.5)
            r3.font.bold = True
            if "عالية" in priority:
                r3.font.color.rgb = RGBColor(185, 28, 28)
            elif "متوسطة" in priority:
                r3.font.color.rgb = RGBColor(180, 83, 9)
            else:
                r3.font.color.rgb = RGBColor(37, 99, 235)

            # Tech Stack
            p4 = row_cells[4].paragraphs[0]
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.paragraph_format.bidi = True
            r4 = p4.add_run(tech)
            r4.font.name = 'Segoe UI'
            r4.font.size = Pt(8)
            r4.font.color.rgb = RGBColor(71, 85, 105)

            # Interactive Checkbox
            p5 = row_cells[5].paragraphs[0]
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p5.paragraph_format.bidi = True
            r5 = p5.add_run("☐")
            r5.font.name = 'Segoe UI Symbol'
            r5.font.size = Pt(13)
            r5.font.color.rgb = RGBColor(100, 116, 139)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Conclusion & Recommendation Section
    doc.add_page_break()
    add_heading_rtl(doc, "🎯 خطة التنفيذ المقترحة والجدول الزمني (Implementation Matrix):", level=1, color_rgb=(15, 23, 42))
    
    plan_text = [
        "1. المرحلة الأولى (Quick Wins - الأسبوع 1 إلى 3): تفعيل ميزات الأولوية العالية في الواتساب والرواتب الذكية والتوقيع الإلكتروني المشفر.",
        "2. المرحلة الثانية (Core Enhancements - الأسبوع 4 إلى 8): إطلاق محرك التداخلات الدوائية المتقدم والتحقق بالوجه وتوليد عقود الصيانة الدورية.",
        "3. المرحلة الثالثة (Scale & Monetization - الأسبوع 9 إلى 12): ربط بوابات الدفع الإلكتروني الليبية (تداول، سداد، ت-باي) وباقة الـ White Label للمشتركين الكبار."
    ]
    for pt in plan_text:
        add_paragraph_rtl(doc, pt, font_size=10.5)

    output_path = r"I:\at\100_Innovative_Ideas_HodoorK_System.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == '__main__':
    try:
        generate_document()
    except Exception as e:
        import traceback
        traceback.print_exc()
        sys.exit(1)
